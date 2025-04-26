import msal
import requests
import webbrowser
import fitz  # PyMuPDF
import re
from io import BytesIO
import openai
import xml.sax.saxutils as xml_utils
import xml.etree.ElementTree as ET
from datetime import datetime
import json
import csv
import os
from fpdf import FPDF
from bs4 import BeautifulSoup
from office365.sharepoint.client_context import ClientContext
from office365.sharepoint.files.file import File
from office365.runtime.auth.client_credential import ClientCredential
from docx import Document  # For handling .docx files
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.service_account import Credentials
from google.oauth2 import service_account
from dotenv import load_dotenv

# Load environment variables from .env file if it exists
load_dotenv()

# --- CONFIG ---
CLIENT_ID = os.environ.get("MS_CLIENT_ID", "")
TENANT_ID = os.environ.get("MS_TENANT_ID", "")
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
SCOPES = ["Group.Read.All", "User.Read"]
GROUP_ID = os.environ.get("MS_GROUP_ID", "")
ACCEPTABLE_LOCATIONS = {"oregon", "portland", "mexico", "or"}
LOCAL_FOLDER = "Applicants"

# Get OpenAI API key from environment variable
openai.api_key = os.environ.get("OPENAI_API_KEY", "")

# SharePoint Configuration
SHAREPOINT_SITE_URL = os.environ.get("SHAREPOINT_SITE_URL", "https://aheadcomputinginc.sharepoint.com/sites/ACHR")
SHAREPOINT_CLIENT_SECRET = os.environ.get("SHAREPOINT_CLIENT_SECRET", "")

# Skip the first 50 threads for debugging purposes
DEBUG_SKIP_THREADS = 0
DEBUG_END_THREAD = 2

class Applicant:
    def __init__(
        self,
        name="--",
        email="--",
        phone="--",
        location="--",
        country="--",  # New attribute
        position_category="--",
        university="--",
        degree="--",
        graduation_year="--",
        years_experience="--",
        notable_companies=None,
        top_skills=None,
        job_intention="--",
        visa_status="--",
        able="--",
        gpt_summary="--",
        subject="--",
        date_sent="--",
        folder_path="--",  # New attribute
        resume="No",  # Attribute to indicate if a resume is present
        desired_job_role="--"  # New attribute for the desired job role
    ):
        self.name = name
        self.email = email
        self.phone = phone
        self.location = location
        self.country = country  # Initialize new attribute
        self.position_category = position_category
        self.university = university
        self.degree = degree
        self.graduation_year = graduation_year
        self.years_experience = years_experience
        self.notable_companies = notable_companies or []
        self.top_skills = top_skills or []
        self.job_intention = job_intention
        self.visa_status = visa_status
        self.able = able
        self.gpt_summary = gpt_summary
        self.subject = subject
        self.date_sent = date_sent
        self.folder_path = folder_path  # Initialize new attribute
        self.resume = resume
        self.desired_job_role = desired_job_role  # Initialize new attribute


class Job:
    keywords = {
        "Design": ["logic design", "physical design"],
        "Verification": ["formal verification", "design verification", "emulation"],
        "Arch": ["architecture", "arch"],
        "Internship": ["internship"],
        "Operations": ["hr"]
    }

def acquire_token_interactively():
    cache_file = "token_cache.json"  # File to store the token cache
    app = msal.PublicClientApplication(CLIENT_ID, authority=AUTHORITY, token_cache=msal.SerializableTokenCache())

    # Load the token cache if it exists
    if os.path.exists(cache_file):
        with open(cache_file, "r") as f:
            app.token_cache.deserialize(f.read())

    # Check if a valid token exists in the cache
    accounts = app.get_accounts()
    if accounts:
        result = app.acquire_token_silent(SCOPES, account=accounts[0])
        if result and "access_token" in result:
            return result["access_token"]

    # If no valid token, prompt the user to log in
    flow = app.initiate_device_flow(scopes=SCOPES)
    if "user_code" not in flow:
        raise ValueError("❌ Failed to create device flow")
    print(f"🔐 Authenticate at: {flow['verification_uri']}\nCode: {flow['user_code']}")
    webbrowser.open(flow['verification_uri'])
    result = app.acquire_token_by_device_flow(flow)

    if "access_token" not in result:
        raise RuntimeError(f"❌ Token error: {result.get('error_description')}")

    # Save the token cache to a file
    with open(cache_file, "w") as f:
        f.write(app.token_cache.serialize())

    return result["access_token"]


def generate_structured_summary(email_text, resume_text):
    

    prompt = (
        "You are a helpful assistant that extracts applicant information for a hiring system.\n"
        "Given the following email and resume, extract the following fields and return them as JSON and if any of the fields are not relevant return --:\n\n"
        "Fields:\n"
        "- full_name (Capitalize First Letter Of Names)\n"
        "- email\n"
        "- phone\n"
        "- location\n"
        "- country (determine the country or region based on the location. If United States Name say USA)\n"
        "- position_category (Design: Logic design, Design: Physical Design, Verification: Formal verification, Verification: Design verification,"
        "  Architecture, Internship, Operations: Human Resources, Other)\n"
        "- university (or universities)\n"
        "- degree\n"
        "- graduation_year (estimate if not clearly stated)\n"
        "- years_experience (estimate)\n"
        "- notable_companies (list)\n"
        "- top_skills (list)\n"
        "- job_intention (e.g. full-time, internship, etc.)\n"
        "- visa_status (if mentioned)\n"
        "- able (yes if their location is in Oregon or Mexico and not looking for an internship, no if their location is somewhere else or looking for internship)\n"
        "- summary (a brief natural language summary)\n"
        "- application (Yes if the email sender is requesting consideration for employment at aheadcomputing. No if otherwise.)\n"
        "- explanation (if the application field is no, explain why and what it is if not an application, if someone on our team has reached out and rejected them say so.)\n"
        "- application_score (scale of 0 to 100 based on the likelyhood that the email sender is requesting employment. 0 being unlikely 100 being most likely)\n"
        "- desired_job_role (Based on their resume, skills, and email content, provide the most likely specific job role/title they want regardless of what we are offering. Be very specific with the job title - for example, don't just say 'engineer', specify what kind like 'FPGA Design Engineer' or 'Machine Learning Engineer'. This should represent their ideal job based on their qualifications.)\n\n"
        "Email:\n"
        f"{email_text}\n\n"
        "Resume:\n"
        f"{resume_text}\n\n"
        "Return ONLY valid JSON with the fields above."
    )
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=700
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        return f"GPT structured summary failed: {e}"


def extract_pdf_text(pdf_bytes):
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)

def extract_docx_text(docx_bytes):
    with BytesIO(docx_bytes) as docx_stream:
        document = Document(docx_stream)
        # Extract text from the main body
        text = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Extract text from headers
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

        # Extract text from footers
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return "\n".join(text)


def display_applicant_info(applicant):
    if all(attr == "--" for attr in [applicant.name, applicant.email, applicant.phone]):
        return
    print(f"\n📧 Subject: {applicant.subject}")
    print(f"Name: {applicant.name}\nEmail: {applicant.email}\nPhone: {applicant.phone}")
    # print(f"Education:\n{applicant.education}\n\nExperience:\n{applicant.experience}")
    print(f"Position Category: {applicant.role}")
    print(f"Location: {applicant.location}")
    print(f"GPT Summary: {applicant.gpt_summary}")
    print(f"Date Sent: {applicant.date_sent}")
    print("-" * 50)


def generate_folder_path(applicant):
    base_folder = "Applicants"

    # Determine the country folder
    country_folder = "USA" if "usa" in str(applicant.country).lower() else (
        "Mexico" if "mexico" in str(applicant.country).lower() else "OtherOrUnknown"
    )

    # Determine the position folder and subfolder
    # Handle position_category as either string or list
    if isinstance(applicant.position_category, list):
        position_category = str(applicant.position_category[0]).lower() if applicant.position_category else "other"
    else:
        position_category = str(applicant.position_category).lower() if applicant.position_category else "other"
    
    if "logic design" in position_category:
        position_folder = "Design/Logic Design"
    elif "physical design" in position_category:
        position_folder = "Design/Physical Design"
    elif "formal verification" in position_category:
        position_folder = "Verification/Formal Verification"
    elif "design verification" in position_category:
        position_folder = "Verification/Design Verification"
    elif "architecture" in position_category:
        position_folder = "Architecture"
    elif "human resources" in position_category:
        position_folder = "Operations/Human Resources"
    elif "internship" in position_category:
        position_folder = "Internships"
    else:
        position_folder = "Other"

    # Format the applicant folder
    if applicant.name != "--" and applicant.date_sent != "--":
        applicant_folder = f"{applicant.name}: {applicant.date_sent}"
    else:
        # Fallback to -- if date is unknown
        fallback_name = applicant.name if applicant.name != "--" else "Unknown"
        fallback_date = "--"
        applicant_folder = f"{fallback_name}: {fallback_date}"

    # Combine all parts into the full folder path
    return f"{base_folder}/{country_folder}/{position_folder}/{applicant_folder}"


def regenerate_full_xml(applicants, filename="applicants.xml"):
    root = ET.Element("Applicants")
    for app in applicants:
        if all(attr == "--" for attr in [app.name, app.email, app.phone]):
            continue
        entry = ET.SubElement(root, "Applicant")
        ET.SubElement(entry, "Name").text = str(app.name)
        ET.SubElement(entry, "Email").text = str(app.email)
        ET.SubElement(entry, "Phone").text = str(app.phone)
        ET.SubElement(entry, "Location").text = str(app.location)
        ET.SubElement(entry, "Country").text = str(app.country)
        ET.SubElement(entry, "PositionCategory").text = str(app.position_category)
        ET.SubElement(entry, "University").text = str(app.university)
        ET.SubElement(entry, "Degree").text = str(app.degree)
        ET.SubElement(entry, "GraduationYear").text = str(app.graduation_year)
        ET.SubElement(entry, "YearsExperience").text = str(app.years_experience)
        ET.SubElement(entry, "JobIntention").text = str(app.job_intention)
        ET.SubElement(entry, "VisaStatus").text = str(app.visa_status)
        ET.SubElement(entry, "Able").text = str(app.able)
        ET.SubElement(entry, "Subject").text = str(app.subject)
        ET.SubElement(entry, "DateSent").text = str(app.date_sent)
        ET.SubElement(entry, "NotableCompanies").text = ", ".join(map(str, app.notable_companies))
        ET.SubElement(entry, "TopSkills").text = ", ".join(map(str, app.top_skills))
        ET.SubElement(entry, "GPTSummary").text = xml_utils.escape(str(app.gpt_summary))
        ET.SubElement(entry, "FolderPath").text = str(app.folder_path)
        ET.SubElement(entry, "Resume").text = str(app.resume)
        ET.SubElement(entry, "DesiredJobRole").text = str(app.desired_job_role)

    tree = ET.ElementTree(root)
    tree.write(filename, encoding="utf-8", xml_declaration=True)


def append_new_applicants(applicants, filename="applicant_bank.xml"):
    # Load existing applicants from XML
    try:
        tree = ET.parse(filename)
        root = tree.getroot()
    except FileNotFoundError:
        root = ET.Element("Applicants")

    # Check for both possible name tag formats (n and Name)
    existing_emails = {app.find("Email").text for app in root.findall("Applicant") if app.find("Email") is not None}
    existing_names = set()
    
    for app in root.findall("Applicant"):
        name_element = app.find("n") if app.find("n") is not None else app.find("Name")
        if name_element is not None and name_element.text is not None:
            existing_names.add(name_element.text)

    for app in applicants:
        if app.email in existing_emails and app.name in existing_names:
            continue  # Skip if applicant already exists by email and name

        entry = ET.SubElement(root, "Applicant")
        # Use the same tag name format as the existing file
        if root.find("Applicant/n") is not None:
            ET.SubElement(entry, "n").text = str(app.name)
        else:
            ET.SubElement(entry, "Name").text = str(app.name)
            
        ET.SubElement(entry, "Email").text = str(app.email)
        ET.SubElement(entry, "Phone").text = str(app.phone)
        ET.SubElement(entry, "Location").text = str(app.location)
        ET.SubElement(entry, "Country").text = str(app.country)
        ET.SubElement(entry, "PositionCategory").text = str(app.position_category)
        ET.SubElement(entry, "University").text = str(app.university)
        ET.SubElement(entry, "Degree").text = str(app.degree)
        ET.SubElement(entry, "GraduationYear").text = str(app.graduation_year)
        ET.SubElement(entry, "YearsExperience").text = str(app.years_experience)
        ET.SubElement(entry, "JobIntention").text = str(app.job_intention)
        ET.SubElement(entry, "VisaStatus").text = str(app.visa_status)
        ET.SubElement(entry, "Able").text = str(app.able)
        ET.SubElement(entry, "Subject").text = str(app.subject)
        ET.SubElement(entry, "DateSent").text = str(app.date_sent)
        ET.SubElement(entry, "NotableCompanies").text = ", ".join(map(str, app.notable_companies))
        ET.SubElement(entry, "TopSkills").text = ", ".join(map(str, app.top_skills))
        ET.SubElement(entry, "GPTSummary").text = xml_utils.escape(str(app.gpt_summary))
        ET.SubElement(entry, "FolderPath").text = str(app.folder_path)
        ET.SubElement(entry, "Resume").text = str(app.resume)
        ET.SubElement(entry, "DesiredJobRole").text = str(app.desired_job_role)

    tree = ET.ElementTree(root)
    tree.write(filename, encoding="utf-8", xml_declaration=True)


def save_single_applicant_to_xml(applicant, filename="applicant_bank.xml"):
    """
    Saves a single applicant to an XML file, checking for duplicates first.
    
    Args:
        applicant: The Applicant object to save
        filename: The XML file to save to (default: applicant_bank.xml)
    """
    try:
        tree = ET.parse(filename)
        root = tree.getroot()
    except FileNotFoundError:
        root = ET.Element("Applicants")

    existing_emails = {app.find("Email").text for app in root.findall("Applicant")}
    existing_names = {app.find("n").text if app.find("n") is not None else app.find("Name").text 
                     for app in root.findall("Applicant") if app.find("n") is not None or app.find("Name") is not None}

    if applicant.email in existing_emails and applicant.name in existing_names:
        return  # Skip if applicant already exists by email and name

    entry = ET.SubElement(root, "Applicant")
    
    # Use the correct tag names based on filename
    if filename == "applicant_bank.xml":
        ET.SubElement(entry, "n").text = str(applicant.name)
    else:
        ET.SubElement(entry, "Name").text = str(applicant.name)
        
    ET.SubElement(entry, "Email").text = str(applicant.email)
    ET.SubElement(entry, "Phone").text = str(applicant.phone)
    ET.SubElement(entry, "Location").text = str(applicant.location)
    ET.SubElement(entry, "Country").text = str(applicant.country)
    ET.SubElement(entry, "PositionCategory").text = str(applicant.position_category)
    ET.SubElement(entry, "University").text = str(applicant.university)
    ET.SubElement(entry, "Degree").text = str(applicant.degree)
    ET.SubElement(entry, "GraduationYear").text = str(applicant.graduation_year)
    ET.SubElement(entry, "YearsExperience").text = str(applicant.years_experience)
    ET.SubElement(entry, "JobIntention").text = str(applicant.job_intention)
    ET.SubElement(entry, "VisaStatus").text = str(applicant.visa_status)
    ET.SubElement(entry, "Able").text = str(applicant.able)
    ET.SubElement(entry, "Subject").text = str(applicant.subject)
    ET.SubElement(entry, "DateSent").text = str(applicant.date_sent)
    ET.SubElement(entry, "NotableCompanies").text = ", ".join(map(str, applicant.notable_companies))
    ET.SubElement(entry, "TopSkills").text = ", ".join(map(str, applicant.top_skills))
    ET.SubElement(entry, "GPTSummary").text = xml_utils.escape(str(applicant.gpt_summary))
    ET.SubElement(entry, "FolderPath").text = str(applicant.folder_path)
    ET.SubElement(entry, "Resume").text = str(applicant.resume)
    ET.SubElement(entry, "DesiredJobRole").text = str(applicant.desired_job_role)

    tree = ET.ElementTree(root)
    tree.write(filename, encoding="utf-8", xml_declaration=True)


def save_applicants_to_csv(applicants, filename="applicants.csv"):
    with open(filename, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)

        # Header row
        writer.writerow([
            "Name", "Email", "Phone", "Location", "Country", "Position Category", "University", "Degree",
            "Graduation Year", "Years Experience", "Job Intention", "Visa Status", "Able",
            "Notable Companies", "Top Skills", "Subject", "Date Sent", "GPT Summary", "Folder Path", "Resume",
            "Desired Job Role"  # Added new column
        ])

        # Data rows
        for app in applicants:
            if all(attr == "--" for attr in [app.name, app.email, app.phone]):
                continue

            writer.writerow([
                app.name,
                app.email,
                app.phone,
                app.location,
                app.country,  # Include country in CSV
                app.position_category,
                app.university,
                app.degree,
                app.graduation_year,
                app.years_experience,
                app.job_intention,
                app.visa_status,
                app.able,
                ", ".join(app.notable_companies),
                ", ".join(app.top_skills),
                app.subject,
                app.date_sent,
                app.gpt_summary.replace("\n", " ").replace("\r", " "),
                app.folder_path,
                app.resume,
                app.desired_job_role  
            ])


def sanitize_text(text):
    if isinstance(text, list):
        # Join list elements into a single string
        text = ", ".join(map(str, text))
    elif not isinstance(text, str):
        # Convert non-string types (e.g., int) to string
        text = str(text)
    return text.encode("ascii", "ignore").decode("ascii")


def clean_html(html_content):
    # Use BeautifulSoup to parse and clean the HTML content
    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text(separator="\n").strip()
    # Reduce excessive newlines to a single newline
    cleaned_text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    # Remove common useless text patterns
    cleaned_text = re.sub(r"You don't often get email from .*?\.\nLearn why this is important", "", cleaned_text, flags=re.IGNORECASE)
    return cleaned_text


# Function to sanitize file names for SharePoint and file system
def sanitize_file_name(name):
    """
    Replace characters that are not allowed in SharePoint or file system file names.
    
    Args:
        name (str): The original file name
        
    Returns:
        str: Sanitized file name with invalid characters replaced by hyphens
    """
    if not isinstance(name, str):
        name = str(name)
        
    invalid_chars = [':', '\\', '/', '*', '?', '"', '<', '>', '|', '#', '{', '}', '%', '~', '&', "'"]
    sanitized = name
    for char in invalid_chars:
        sanitized = sanitized.replace(char, '-')
    return sanitized


def generate_applicant_folder(applicant, email_text, attachments, attachment_metadata):
    # Create the folder path
    folder_path = applicant.folder_path
    os.makedirs(folder_path, exist_ok=True)  # Ensure the folder exists

    # Save email-attached PDFs with original filenames or fallback
    for i, (attachment, filename) in enumerate(zip(attachments, attachment_metadata)):
        sanitized_filename = sanitize_file_name(filename)  # Sanitize the filename
        attachment_path = os.path.join(folder_path, sanitized_filename)
        with open(attachment_path, "wb") as f:
            f.write(attachment)  # Save the attachment content

    # Overwrite the PDF of the email text
    email_pdf_path = os.path.join(folder_path, "email_text.pdf")
    email_pdf = FPDF()
    email_pdf.add_page()
    email_pdf.set_font("Arial", size=12)
    email_text = clean_html(email_text)  # Ensure HTML is cleaned
    email_text = sanitize_text(email_text)
    email_pdf.multi_cell(0, 10, email_text)
    email_pdf.output(email_pdf_path)  # Always overwrite existing files

    # Overwrite the PDF of the parsed summary info
    summary_pdf_path = os.path.join(folder_path, "summary.pdf")
    summary_pdf = FPDF()
    summary_pdf.add_page()
    summary_pdf.set_font("Arial", size=12)

    # Add GPT summary at the top
    sanitized_summary = sanitize_text(applicant.gpt_summary)
    summary_pdf.multi_cell(0, 10, f"Summary:\n{sanitized_summary}\n\n")

    # Sanitize and write other applicant details
    summary_pdf.multi_cell(0, 10, f"Name: {sanitize_text(applicant.name)}")
    summary_pdf.multi_cell(0, 10, f"Email: {sanitize_text(applicant.email)}")
    summary_pdf.multi_cell(0, 10, f"Phone: {sanitize_text(applicant.phone)}")
    summary_pdf.multi_cell(0, 10, f"Location: {sanitize_text(applicant.location)}")
    summary_pdf.multi_cell(0, 10, f"Country: {sanitize_text(applicant.country)}")
    summary_pdf.multi_cell(0, 10, f"Position Category: {sanitize_text(applicant.position_category)}")
    summary_pdf.multi_cell(0, 10, f"Desired Job Role: {sanitize_text(applicant.desired_job_role)}")
    summary_pdf.multi_cell(0, 10, f"University: {sanitize_text(applicant.university)}")
    summary_pdf.multi_cell(0, 10, f"Degree: {sanitize_text(applicant.degree)}")
    summary_pdf.multi_cell(0, 10, f"Graduation Year: {sanitize_text(applicant.graduation_year)}")
    summary_pdf.multi_cell(0, 10, f"Years Experience: {sanitize_text(applicant.years_experience)}")
    summary_pdf.multi_cell(0, 10, f"Notable Companies: {sanitize_text(applicant.notable_companies)}")
    summary_pdf.multi_cell(0, 10, f"Top Skills: {sanitize_text(applicant.top_skills)}")
    summary_pdf.multi_cell(0, 10, f"Job Intention: {sanitize_text(applicant.job_intention)}")
    summary_pdf.multi_cell(0, 10, f"Visa Status: {sanitize_text(applicant.visa_status)}")
    summary_pdf.multi_cell(0, 10, f"Able: {sanitize_text(applicant.able)}")
    summary_pdf.multi_cell(0, 10, f"Subject: {sanitize_text(applicant.subject)}")
    summary_pdf.multi_cell(0, 10, f"Date Sent: {sanitize_text(applicant.date_sent)}")
    summary_pdf.multi_cell(0, 10, f"Resume: {sanitize_text(applicant.resume)}")

    summary_pdf.output(summary_pdf_path)  # Always overwrite existing files

# Function to detect Google Drive links
def detect_google_drive_links(email_text):
    drive_link_pattern = r"https://drive\.google\.com/file/d/[\w-]+/view\?usp=sharing"
    return re.findall(drive_link_pattern, email_text)


# Function to retrieve attachment text and Google Drive links
def retrieve_attachment_text(messages, headers, thread_id):
    google_drive_links = []
    attachment_data = []
    attachment_metadata = []  # To store metadata for filenames
    resume_text = ""  # Initialize resume_text to accumulate text from all attachments
    

    for msg in messages:
        email_text = msg.get("body", {}).get("content", "--")
        email_text = clean_html(email_text)  # Clean the email text to remove HTML and noise

        # Check for Google Drive links
        google_drive_links.extend(detect_google_drive_links(email_text))

        attachments_url = f"https://graph.microsoft.com/v1.0/groups/{GROUP_ID}/threads/{thread_id}/posts/{msg['id']}/attachments"
        attachments = requests.get(attachments_url, headers=headers).json().get("value", [])

        for att in attachments:
            if att.get("contentType") == "application/pdf" and att.get("id"):
                att_id = att["id"]
                att_url = f"https://graph.microsoft.com/v1.0/groups/{GROUP_ID}/threads/{thread_id}/posts/{msg['id']}/attachments/{att_id}/$value"
                file_data = requests.get(att_url, headers=headers).content
                attachment_data.append(file_data)
                attachment_metadata.append(att.get("name", f"attachment_{len(attachment_data)}.pdf"))  # Save original filename or fallback
                resume_text += extract_pdf_text(file_data) + "\n"  # Append extracted text to resume_text

            elif att.get("contentType") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and att.get("id"):
                att_id = att["id"]
                att_url = f"https://graph.microsoft.com/v1.0/groups/{GROUP_ID}/threads/{thread_id}/posts/{msg['id']}/attachments/{att_id}/$value"
                file_data = requests.get(att_url, headers=headers).content
                attachment_data.append(file_data)
                attachment_metadata.append(att.get("name", f"attachment_{len(attachment_data)}.docx"))  # Save original filename or fallback
                resume_text += extract_docx_text(file_data) + "\n"  # Append extracted text to resume_text
    
    # Process Google Drive links if no attachments were found or as additional documents
    if google_drive_links:
        for link in google_drive_links:
            try:
                # Extract the file ID from the link
                file_id = re.search(r"https://drive\.google\.com/file/d/([\w-]+)/view", link).group(1)
                
                # Authenticate using the service account credentials
                credentials = Credentials.from_service_account_file("service_account.json", 
                                                                   scopes=["https://www.googleapis.com/auth/drive"])
                service = build("drive", "v3", credentials=credentials)
                
                # Get file metadata to determine name and type
                file_metadata = service.files().get(fileId=file_id, fields="mimeType,name").execute()
                mime_type = file_metadata.get("mimeType")
                file_name = file_metadata.get("name", f"{file_id}")
                
                # If no extension in filename, add appropriate extension based on mime type
                if "." not in file_name:
                    if mime_type == "application/pdf":
                        file_name += ".pdf"
                    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        file_name += ".docx"
                
                # Request the file content
                request = service.files().get_media(fileId=file_id)
                file_data = BytesIO()
                downloader = MediaIoBaseDownload(file_data, request)
                
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                    print(f"Google Drive download progress: {int(status.progress() * 100)}%")
                
                # Add to attachments
                file_data.seek(0)
                attachment_data.append(file_data.read())
                attachment_metadata.append(file_name)
                
                # Extract text based on mime type
                file_data.seek(0)
                if mime_type == "application/pdf":
                    extracted_text = extract_pdf_text(file_data.read())
                elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    extracted_text = extract_docx_text(file_data.read())
                else:
                    extracted_text = f"Unsupported file type: {mime_type}"
                
                resume_text += extracted_text + "\n"
                print(f"Successfully processed Google Drive file: {file_name}")
                
            except Exception as e:
                print(f"Error processing Google Drive link: {e}")
                
    # If no resume text was found, mark it as "--"
    if not resume_text.strip():
        resume_text = "--"

    return resume_text, attachment_data, attachment_metadata


# Updated logic to handle resumes and categorize applicants dynamically
def process_email_chain(messages, headers, thread_id):
    email_text_combined = ""

    for msg in messages:
        email_text = msg.get("body", {}).get("content", "--")
        email_text = clean_html(email_text)  # Clean the email text to remove HTML and noise
        email_text_combined += email_text + "\n\n"  # Add spacing for readability

    resume_text, attachment_data, attachment_metadata = retrieve_attachment_text(messages, headers, thread_id)

    # Use GPT to determine if the email chain is from an applicant
    gpt_response = generate_structured_summary(email_text_combined, resume_text)  # Ensure resume_text is passed
    try:
        parsed = json.loads(gpt_response)
        is_applicant = parsed.get("application", "no").lower() == "yes"
    except:
        parsed = {}
        is_applicant = False

    # Ensure 'name' key exists in the parsed dictionary before accessing it
    if parsed.get("full_name") == "--":
        is_applicant = False
   

    return is_applicant, parsed, resume_text, email_text_combined, attachment_data, attachment_metadata


def format_date_from_message(messages):
    """
    Extracts and formats the creation date from the first message in the list.

    Args:
        messages (list): A list of message dictionaries.

    Returns:
        str: The formatted date in 'YYYY-MM-DD' format, or '--' if parsing fails.
    """
    raw_date = messages[0].get("createdDateTime", "--")
    try:
        parsed_date = datetime.strptime(raw_date, "%Y-%m-%dT%H:%M:%SZ")
        formatted_date = parsed_date.strftime("%m-%d-%Y")
    except:
        formatted_date = "--"

    return formatted_date



def upload_folder_to_sharepoint(site_url, client_id, client_secret, folder_path, sharepoint_folder):
    """
    Uploads a local folder and its contents to a SharePoint folder using REST API with token authentication.
    This uses the token we already have in token_user_cache.json from our successful test_sharepoint.py script.
    """
    import logging
    import msal
    import time
    import webbrowser
    import requests
    import json
    from urllib.parse import urlparse, quote
    from time import sleep

    # Configure logging
    logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")

    # Parse the site URL to get the host and site name
    parsed_url = urlparse(site_url)
    sharepoint_host = parsed_url.netloc
    site_path = '/'.join(parsed_url.path.split('/')[2:])  # Get site name after /sites/
    
    print(f"🔄 Starting SharePoint upload process...")
    print(f"Site URL: {site_url}")
    print(f"Target folder: {sharepoint_folder}")
    print(f"Local folder: {folder_path}")
    
    # Get the tenant ID from the global variable
    tenant_id = TENANT_ID
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    
    # SharePoint resource ID for correct audience
    sharepoint_resource_id = "00000003-0000-0ff1-ce00-000000000000"
    
    # These are the correct scopes for SharePoint access
    scopes = [f"{sharepoint_resource_id}/.default"]
    
    # Load cache from file if it exists
    cache_file = "token_user_cache.json"
    cache = msal.SerializableTokenCache()
    
    if os.path.exists(cache_file):
        with open(cache_file, "r") as f:
            cache.deserialize(f.read())
    
    # Create public client application with token cache
    app = msal.PublicClientApplication(
        client_id=client_id,
        authority=authority,
        token_cache=cache
    )
    
    # Get token - try silent auth first, then interactive if needed
    accounts = app.get_accounts()
    if accounts:
        logging.info(f"Found {len(accounts)} account(s) in cache. Attempting silent auth.")
        result = app.acquire_token_silent(scopes, account=accounts[0])
        if result and "access_token" in result:
            token = result["access_token"]
            logging.info("Successfully obtained token from cache.")
            # Log token length and first/last few chars for debugging
            logging.debug(f"Token length: {len(token)}")
            logging.debug(f"Token preview: {token[:10]}...{token[-10:]}")
            # Save cache
            with open(cache_file, "w") as f:
                f.write(cache.serialize())
        else:
            logging.warning("No valid token in cache or token expired.")
            token = None
    else:
        logging.warning("No account in cache.")
        token = None
    
    # If silent auth fails, fall back to device code flow
    if token is None:
        logging.info("Starting device code flow for authentication...")
        flow = app.initiate_device_flow(scopes=scopes)
        
        if "user_code" not in flow:
            error_msg = f"Failed to start device flow: {flow.get('error_description', flow.get('error'))}"
            logging.error(error_msg)
            print(f"❌ Authentication error: {error_msg}")
            return False
        
        # Display instructions for the user
        print("\n" + "="*50)
        print("🔐 AUTHENTICATION REQUIRED FOR SHAREPOINT UPLOAD 🔐")
        print("="*50)
        print("\nTo upload to SharePoint, you need to sign in with your Microsoft account.")
        print(f"\n1. Go to: {flow['verification_uri']}")
        print(f"2. Enter this code: {flow['user_code']}")
        print(f"\n3. Sign in with your Microsoft account that has access to SharePoint.")
        print("="*50 + "\n")
        
        # Open the browser automatically
        try:
            webbrowser.open(flow["verification_uri"])
        except Exception as e:
            logging.warning(f"Could not open browser automatically: {e}")
        
        # Wait for the user to authenticate
        result = app.acquire_token_by_device_flow(flow)
        
        if "access_token" not in result:
            error = result.get("error", "unknown")
            error_desc = result.get("error_description", "No description")
            error_msg = f"Failed to acquire token. Error: {error}. Description: {error_desc}"
            logging.error(error_msg)
            print(f"❌ Authentication error: {error_msg}")
            return False
            
        token = result["access_token"]
        logging.info("✅ Successfully obtained access token!")
        logging.debug(f"Token length: {len(token)}")
        logging.debug(f"Token preview: {token[:10]}...{token[-10:]}")
        
        # Save token cache
        with open(cache_file, "w") as f:
            f.write(cache.serialize())
    
    # Prepare headers for all REST API calls
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json;odata=verbose",
        "Content-Type": "application/json;odata=verbose"
    }
    
    # Test connection by getting web info
    try:
        test_url = f"{site_url}/_api/web"
        response = requests.get(test_url, headers=headers)
        if response.status_code == 200:
            web_info = response.json()
            web_title = web_info.get('d', {}).get('Title', 'Unknown')
            print(f"\n✅ Successfully connected to SharePoint site: {web_title}")
            logging.info(f"Successfully connected to SharePoint site: {web_title}")
        else:
            print(f"\n❌ Connection test failed: {response.status_code} {response.text}")
            logging.error(f"Connection test failed: {response.status_code} {response.text}")
            return False
    except Exception as e:
        print(f"\n❌ Connection test error: {e}")
        logging.error(f"Connection test error: {e}")
        return False
    
    # Get the form digest value (required for POST operations)
    try:
        digest_url = f"{site_url}/_api/contextinfo"
        response = requests.post(digest_url, headers=headers)
        if response.status_code != 200:
            print(f"\n❌ Failed to get form digest: {response.status_code} {response.text}")
            logging.error(f"Failed to get form digest: {response.status_code} {response.text}")
            return False
        
        form_digest = response.json()['d']['GetContextWebInformation']['FormDigestValue']
        # Add the form digest to headers for POST requests
        post_headers = headers.copy()
        post_headers["X-RequestDigest"] = form_digest
    except Exception as e:
        print(f"\n❌ Form digest error: {e}")
        logging.error(f"Form digest error: {e}")
        return False

    # Function to sanitize folder names for SharePoint
    def sanitize_folder_name(name):
        # Replace characters that are not allowed in SharePoint folder names
        invalid_chars = [':', '\\', '/', '*', '?', '"', '<', '>', '|', '#', '{', '}', '%', '~', '&']
        sanitized = name
        for char in invalid_chars:
            sanitized = sanitized.replace(char, '-')
        return sanitized
        
    # Function to create a folder using REST API
    def ensure_folder_exists(folder_path):
        # Split the path into parts
        folder_parts = folder_path.strip('/').split('/')
        current_path = ""
        
        # Process each part of the path, sanitizing each folder name
        for part in folder_parts:
            if not part:
                continue
            
            # Sanitize the folder name
            sanitized_part = sanitize_folder_name(part)
            
            if current_path:
                current_path += f"/{sanitized_part}"
            else:
                current_path = sanitized_part
                
            # Check if folder exists
            folder_url = f"{site_url}/_api/web/GetFolderByServerRelativeUrl('{quote(current_path)}')"
            try:
                response = requests.get(folder_url, headers=headers)
                if response.status_code == 200:
                    logging.debug(f"Folder exists: {current_path}")
                    continue
            except Exception:
                pass  # If error, we'll try to create the folder
                
            # Create the folder if it doesn't exist
            parent_path = '/'.join(current_path.split('/')[:-1])
            folder_name = current_path.split('/')[-1]
            
            if not parent_path:
                # Creating folder at root
                create_url = f"{site_url}/_api/web/folders"
                data = {"__metadata": {"type": "SP.Folder"}, "ServerRelativeUrl": current_path}
            else:
                # Creating subfolder
                create_url = f"{site_url}/_api/web/GetFolderByServerRelativeUrl('{quote(parent_path)}')/folders"
                data = {"__metadata": {"type": "SP.Folder"}, "ServerRelativeUrl": f"{parent_path}/{folder_name}"}
                
            try:
                response = requests.post(
                    create_url,
                    headers=post_headers,
                    data=json.dumps(data)
                )
                
                if response.status_code in (200, 201):
                    logging.info(f"Created folder: {current_path}")
                else:
                    logging.error(f"Failed to create folder {current_path}: {response.status_code} {response.text}")
                    return False
            except Exception as e:
                logging.error(f"Error creating folder {current_path}: {e}")
                return False
                
        return True
    
    # Function to map a local folder path to a sanitized SharePoint folder path
    def map_folder_path(local_folder_path, sharepoint_base_folder, base_local_folder):
        # Get the relative path from the base folder
        relative_path = os.path.relpath(local_folder_path, base_local_folder)
        if relative_path == '.':
            return sharepoint_base_folder
            
        # Split the path into parts and sanitize each folder name
        path_parts = relative_path.replace('\\', '/').split('/')
        sanitized_parts = [sanitize_folder_name(part) for part in path_parts]
        
        # Combine with the SharePoint base folder
        return f"{sharepoint_base_folder}/{'/'.join(sanitized_parts)}"
    
    # Function to upload a file using REST API
    def upload_file(file_path, target_folder):
        file_name = os.path.basename(file_path)
        sanitized_file_name = sanitize_file_name(file_name)  # Sanitize the file name
        encoded_name = quote(sanitized_file_name)
        
        # Start upload session
        upload_url = f"{site_url}/_api/web/GetFolderByServerRelativeUrl('{quote(target_folder)}')/Files/add(url='{encoded_name}',overwrite=true)"
        
        with open(file_path, 'rb') as file_content:
            file_data = file_content.read()
            
        # Add content-length header
        upload_headers = post_headers.copy()
        upload_headers["Content-Length"] = str(len(file_data))
        
        try:
            response = requests.post(
                upload_url,
                headers=upload_headers,
                data=file_data
            )
            
            if response.status_code in (200, 201):
                return True
            else:
                logging.error(f"Failed to upload file {file_name}: {response.status_code} {response.text}")
                return False
        except Exception as e:
            logging.error(f"Error uploading file {file_name}: {e}")
            return False
    
    # Now proceed with uploading files
    total_files = sum([len(files) for _, _, files in os.walk(folder_path)])
    uploaded_files = 0
    failed_files = 0
    
    print(f"\nStarting upload of {total_files} files to SharePoint...")
    
    # Create a mapping of local folders to sanitized SharePoint folders
    folder_mapping = {}
    
    # Process the folders and files
    for root, dirs, files in os.walk(folder_path):
        # Get the sanitized SharePoint folder path
        sharepoint_relative_folder = map_folder_path(root, sharepoint_folder, folder_path)
        folder_mapping[root] = sharepoint_relative_folder
        
        # Log what we're processing
        relative_path = os.path.relpath(root, folder_path)
        if relative_path != '.':
            print(f"Processing folder: {relative_path}...")
            print(f"  → SharePoint path: {sharepoint_relative_folder}")
        
        # Ensure the target folder exists
        if not ensure_folder_exists(sharepoint_relative_folder):
            print(f"❌ Failed to create folder: {sharepoint_relative_folder}")
            continue
        
        # Upload each file in the current folder
        for file_name in files:
            local_file_path = os.path.join(root, file_name)
            
            try:
                # Upload the file
                if upload_file(local_file_path, sharepoint_relative_folder):
                    logging.info(f"Successfully uploaded file: {file_name} to {sharepoint_relative_folder}")
                    uploaded_files += 1
                else:
                    failed_files += 1
                    print(f"❌ Failed to upload file: {file_name}")
                
                # Print progress
                progress = (uploaded_files + failed_files) / total_files * 100
                print(f"Progress: {uploaded_files}/{total_files} files uploaded ({progress:.1f}%)", end='\r')
                
                # Add a small delay to avoid throttling
                sleep(0.05)
                
            except Exception as e:
                logging.error(f"Failed to upload file: {local_file_path} to {sharepoint_relative_folder}. Error: {e}")
                failed_files += 1
                print(f"❌ Failed to upload file: {file_name}")
    
    print(f"\n✅ SharePoint upload completed: {uploaded_files} files uploaded, {failed_files} files failed.")
    logging.info(f"SharePoint upload process completed: {uploaded_files} successful, {failed_files} failed.")
    
    return uploaded_files > 0


def download_google_drive_file(file_id, destination_path):
    """Download a file from Google Drive given its file ID."""
    try:
        # Authenticate using the service account
        credentials = service_account.Credentials.from_service_account_file(
            'service_account.json',
            scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=credentials)

        # Request the file
        request = service.files().get_media(fileId=file_id)
        with open(destination_path, 'wb') as file:
            downloader = MediaIoBaseDownload(file, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
                print(f"Download progress: {int(status.progress() * 100)}%")
        print(f"File downloaded to {destination_path}")
    except Exception as e:
        print(f"Failed to download file: {e}")


def handle_google_drive_link(link, destination_folder):
    """Process a Google Drive link and download the file to the destination folder."""
    match = re.search(r"https://drive.google.com/file/d/(.*?)/", link)
    if match:
        file_id = match.group(1)
        destination_path = os.path.join(destination_folder, f"{file_id}.pdf")
        download_google_drive_file(file_id, destination_path)
        return destination_path
    else:
        print("Invalid Google Drive link format.")
        return None

# Example usage in the applicant processing pipeline
def process_resume_link(resume_link, applicant_folder):
    """Process a resume link, handling Google Drive links if necessary."""
    if "drive.google.com" in resume_link:
        return handle_google_drive_link(resume_link, applicant_folder)
    else:
        print("Non-Google Drive links are not yet supported.")
        return None



def get_group_emails(token, top_n=400):
    count = 0
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    threads_url = f"https://graph.microsoft.com/v1.0/groups/{GROUP_ID}/threads?$top={top_n}"
    threads = requests.get(threads_url, headers=headers).json().get("value", [])
    applicants = []

    thread_counter = 1

    for thread in threads:
        if thread_counter <= DEBUG_SKIP_THREADS:
            thread_counter += 1
            continue

        subject = thread.get("topic", "--")
        print(f"🔄 Processing thread {thread_counter}: {subject}")
        thread_counter += 1
        thread_id = thread["id"]
        msg_url = f"https://graph.microsoft.com/v1.0/groups/{GROUP_ID}/threads/{thread_id}/posts"
        messages = requests.get(msg_url, headers=headers).json().get("value", [])
        if not messages:
            continue

        formatted_date = format_date_from_message(messages)

        is_applicant, parsed, resume_text, email_text, attachment_data, attachment_metadata = process_email_chain(messages, headers, thread_id)

        print(f"Applicant score: {parsed.get('application_score', '--')} Name: {parsed.get('full_name', '--')} Email: {parsed.get('email', '--')}")

        if not is_applicant:
            print(f"❌ Skipping non-applicant thread: {subject} {formatted_date} REASON: {parsed.get('explanation', 'not an applicant')}\n")
            continue

        applicant = Applicant(
            name=parsed.get("full_name", "--"),
            email=parsed.get("email", "--"),
            phone=parsed.get("phone", "--"),
            location=parsed.get("location", "--"),
            country=parsed.get("country", "--"),
            position_category=parsed.get("position_category", "--"),
            university=parsed.get("university", "--"),
            degree=parsed.get("degree", "--"),
            graduation_year=parsed.get("graduation_year", "--"),
            years_experience=parsed.get("years_experience", "--"),
            notable_companies=parsed.get("notable_companies", []),
            top_skills=parsed.get("top_skills", []),
            job_intention=parsed.get("job_intention", "--"),
            visa_status=parsed.get("visa_status", "--"),
            able=parsed.get("able", "--"),
            gpt_summary=parsed.get("summary", "--"),
            subject=subject,
            date_sent=formatted_date,
            folder_path=generate_folder_path(Applicant(
                name=parsed.get("full_name", "--"),
                location=parsed.get("location", "--"),
                country=parsed.get("country", "--"),
                position_category=parsed.get("position_category", "--"),
                date_sent=formatted_date
            )),
            resume="Yes" if resume_text != "--" else "No",
            desired_job_role=parsed.get("desired_job_role", "--")
        )
        
        generate_applicant_folder(applicant, email_text, attachment_data, attachment_metadata)        
        
        applicants.append(applicant)
        
        # Save this applicant to both XML files immediately after processing
        save_single_applicant_to_xml(applicant, "applicant_bank.xml")
        
        # Also update the current applicants.xml file with all processed applicants so far
        #regenerate_full_xml(applicants)
        
        count += 1
        print(f"*************Applicant {count} processed: {applicant.name}***************\n")

        if thread_counter >= DEBUG_END_THREAD:
            break

    # These are now redundant since we're updating the files after each applicant,
    # but we'll keep them as a final update to ensure everything is saved
    #append_new_applicants(applicants)
    regenerate_full_xml(applicants)


# Example usage for uploading to SharePoint
if __name__ == "__main__":
    print("🔄 Authenticating...")
    token = acquire_token_interactively()
    get_group_emails(token)
    print("🔄 Uploading folder to SharePoint...")
    upload_folder_to_sharepoint(SHAREPOINT_SITE_URL, CLIENT_ID, SHAREPOINT_CLIENT_SECRET, LOCAL_FOLDER, SHAREPOINT_FOLDER)

